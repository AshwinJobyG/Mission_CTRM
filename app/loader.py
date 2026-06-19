"""Read *any* files from a folder into plain text.

Text-like files are read directly; PDFs and DOCX are extracted via optional
libraries when available. Unsupported/binary files are skipped (and reported).
"""

from __future__ import annotations

import csv
import io
import json
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

# Extensions we treat as UTF-8 text and read verbatim.
TEXT_EXTENSIONS = {
    ".txt", ".md", ".markdown", ".rst", ".log",
    ".py", ".js", ".ts", ".java", ".go", ".rb", ".rs", ".c", ".cpp", ".h",
    ".cs", ".php", ".sh", ".bash", ".sql", ".ini", ".cfg", ".conf",
    ".yaml", ".yml", ".toml", ".env", ".html", ".htm", ".xml", ".css",
}
JSON_EXTENSIONS = {".json", ".jsonl", ".ndjson"}
CSV_EXTENSIONS = {".csv", ".tsv"}
PDF_EXTENSIONS = {".pdf"}
DOCX_EXTENSIONS = {".docx"}

# Skip obviously binary/media files outright.
SKIP_EXTENSIONS = {
    ".png", ".jpg", ".jpeg", ".gif", ".bmp", ".svg", ".ico", ".webp",
    ".zip", ".tar", ".gz", ".7z", ".rar", ".exe", ".dll", ".so", ".bin",
    ".mp3", ".mp4", ".mov", ".avi", ".wav", ".ttf", ".woff", ".woff2",
}


@dataclass
class Document:
    """A single source file loaded into text."""

    path: str          # absolute path
    name: str          # file name
    source: str        # path relative to the scanned folder
    text: str
    ext: str


def _read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _read_json(path: Path) -> str:
    raw = path.read_text(encoding="utf-8", errors="ignore")
    try:
        obj = json.loads(raw)
        return json.dumps(obj, indent=2, ensure_ascii=False)
    except Exception:
        # JSONL / malformed — fall back to raw text.
        return raw


def _read_csv(path: Path) -> str:
    delimiter = "\t" if path.suffix.lower() == ".tsv" else ","
    rows: list[str] = []
    with path.open("r", encoding="utf-8", errors="ignore", newline="") as fh:
        reader = csv.reader(fh, delimiter=delimiter)
        for row in reader:
            rows.append(" | ".join(cell.strip() for cell in row))
    return "\n".join(rows)


def _read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except Exception as exc:  # pragma: no cover
        raise RuntimeError("pypdf is required to read PDF files (pip install pypdf)") from exc
    reader = PdfReader(str(path))
    parts: list[str] = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")
    return "\n".join(parts)


def _read_docx(path: Path) -> str:
    try:
        import docx  # python-docx
    except Exception as exc:  # pragma: no cover
        raise RuntimeError("python-docx is required to read .docx files (pip install python-docx)") from exc
    document = docx.Document(str(path))
    return "\n".join(p.text for p in document.paragraphs)


def _extract(path: Path) -> str | None:
    ext = path.suffix.lower()
    if ext in SKIP_EXTENSIONS:
        return None
    if ext in PDF_EXTENSIONS:
        return _read_pdf(path)
    if ext in DOCX_EXTENSIONS:
        return _read_docx(path)
    if ext in JSON_EXTENSIONS:
        return _read_json(path)
    if ext in CSV_EXTENSIONS:
        return _read_csv(path)
    if ext in TEXT_EXTENSIONS or ext == "":
        return _read_text(path)
    # Unknown extension: try as text, give up if it looks binary.
    try:
        text = _read_text(path)
        if "\x00" in text:
            return None
        return text
    except Exception:
        return None


def iter_files(folder: str | Path, recursive: bool = True) -> Iterable[Path]:
    base = Path(folder).expanduser()
    pattern = "**/*" if recursive else "*"
    for p in sorted(base.glob(pattern)):
        if p.is_file():
            yield p


def load_folder(folder: str | Path, recursive: bool = True) -> tuple[list[Document], list[str]]:
    """Load every readable file under ``folder``.

    Returns ``(documents, skipped)`` where ``skipped`` lists files that were
    binary/unsupported or failed to parse.
    """
    base = Path(folder).expanduser()
    if not base.exists():
        raise FileNotFoundError(f"Folder not found: {base}")

    docs: list[Document] = []
    skipped: list[str] = []
    for path in iter_files(base, recursive=recursive):
        try:
            text = _extract(path)
        except Exception as exc:
            skipped.append(f"{path} ({exc})")
            continue
        if not text or not text.strip():
            skipped.append(str(path))
            continue
        rel = str(path.relative_to(base)) if path.is_relative_to(base) else path.name
        docs.append(
            Document(
                path=str(path.resolve()),
                name=path.name,
                source=rel,
                text=text,
                ext=path.suffix.lower(),
            )
        )
    return docs, skipped
